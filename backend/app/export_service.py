"""
Export Service for YouTube Summaries
Handles PDF, DOCX, and Markdown export functionality
"""
import io
import markdown
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
import logging

logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        
    def export_to_markdown(self, session_data: Dict[str, Any]) -> str:
        """Export session to Markdown format"""
        try:
            markdown_content = f"""# YouTube Video Summary

## Video Information
- **Title:** {session_data.get('video_title', 'Unknown')}
- **URL:** {session_data.get('video_url', '')}
- **Duration:** {self._format_duration(session_data.get('video_duration', 0))}
- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Short Summary
{session_data.get('short_summary', 'No summary available')}

## Detailed Summary
{session_data.get('detailed_summary', 'No detailed summary available')}

## Transcript
```
{session_data.get('transcript', 'No transcript available')}
```

## Chat History
"""
            
            chat_history = session_data.get('chat_history', [])
            if chat_history:
                for i, message in enumerate(chat_history):
                    role = message.get('role', 'unknown').title()
                    content = message.get('content', '')
                    timestamp = message.get('timestamp', '')
                    
                    markdown_content += f"\n### {role} ({timestamp})\n{content}\n"
            else:
                markdown_content += "\nNo chat history available.\n"
            
            return markdown_content
            
        except Exception as e:
            logger.error(f"Error exporting to markdown: {e}")
            raise
    
    def export_to_docx(self, session_data: Dict[str, Any]) -> io.BytesIO:
        """Export session to DOCX format"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('YouTube Video Summary', 0)
            title.alignment = 1  # Center alignment
            
            # Video Information
            doc.add_heading('Video Information', level=1)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ['Title', session_data.get('video_title', 'Unknown')],
                ['URL', session_data.get('video_url', '')],
                ['Duration', self._format_duration(session_data.get('video_duration', 0))],
                ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = value
            
            doc.add_page_break()
            
            # Short Summary
            doc.add_heading('Short Summary', level=1)
            doc.add_paragraph(session_data.get('short_summary', 'No summary available'))
            
            # Detailed Summary
            doc.add_heading('Detailed Summary', level=1)
            doc.add_paragraph(session_data.get('detailed_summary', 'No detailed summary available'))
            
            doc.add_page_break()
            
            # Chat History
            doc.add_heading('Chat History', level=1)
            chat_history = session_data.get('chat_history', [])
            
            if chat_history:
                for message in chat_history:
                    role = message.get('role', 'unknown').title()
                    content = message.get('content', '')
                    timestamp = message.get('timestamp', '')
                    
                    # Add role and timestamp
                    role_para = doc.add_paragraph()
                    role_run = role_para.add_run(f"{role} ({timestamp})")
                    role_run.bold = True
                    
                    # Add content
                    doc.add_paragraph(content)
                    doc.add_paragraph()  # Add space
            else:
                doc.add_paragraph('No chat history available.')
            
            doc.add_page_break()
            
            # Transcript
            doc.add_heading('Full Transcript', level=1)
            transcript_para = doc.add_paragraph(session_data.get('transcript', 'No transcript available'))
            transcript_para.style = 'Normal'
            
            # Save to BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise
    
    def export_to_pdf(self, session_data: Dict[str, Any]) -> io.BytesIO:
        """Export session to PDF format"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
            
            # Define styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1,  # Center
                textColor=HexColor('#2563eb')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=self.styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=12,
                textColor=HexColor('#1f2937')
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            )
            
            chat_role_style = ParagraphStyle(
                'ChatRole',
                parent=self.styles['Normal'],
                fontSize=12,
                spaceBefore=10,
                spaceAfter=5,
                textColor=HexColor('#059669')
            )
            
            # Build content
            content = []
            
            # Title
            content.append(Paragraph("YouTube Video Summary", title_style))
            content.append(Spacer(1, 20))
            
            # Video Information
            content.append(Paragraph("Video Information", heading_style))
            
            video_info = f"""
            <b>Title:</b> {session_data.get('video_title', 'Unknown')}<br/>
            <b>URL:</b> {session_data.get('video_url', '')}<br/>
            <b>Duration:</b> {self._format_duration(session_data.get('video_duration', 0))}<br/>
            <b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            """
            content.append(Paragraph(video_info, body_style))
            content.append(Spacer(1, 20))
            
            # Short Summary
            content.append(Paragraph("Short Summary", heading_style))
            content.append(Paragraph(session_data.get('short_summary', 'No summary available'), body_style))
            content.append(Spacer(1, 20))
            
            # Detailed Summary
            content.append(Paragraph("Detailed Summary", heading_style))
            content.append(Paragraph(session_data.get('detailed_summary', 'No detailed summary available'), body_style))
            content.append(PageBreak())
            
            # Chat History
            content.append(Paragraph("Chat History", heading_style))
            chat_history = session_data.get('chat_history', [])
            
            if chat_history:
                for message in chat_history:
                    role = message.get('role', 'unknown').title()
                    content_text = message.get('content', '')
                    timestamp = message.get('timestamp', '')
                    
                    content.append(Paragraph(f"<b>{role}</b> ({timestamp})", chat_role_style))
                    content.append(Paragraph(content_text, body_style))
                    content.append(Spacer(1, 10))
            else:
                content.append(Paragraph('No chat history available.', body_style))
            
            content.append(PageBreak())
            
            # Transcript
            content.append(Paragraph("Full Transcript", heading_style))
            transcript_text = session_data.get('transcript', 'No transcript available')
            # Split long transcript into paragraphs for better PDF formatting
            transcript_paragraphs = transcript_text.split('\n\n')
            for para in transcript_paragraphs:
                if para.strip():
                    content.append(Paragraph(para, body_style))
            
            # Build PDF
            doc.build(content)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    def _format_duration(self, seconds: int) -> str:
        """Format duration in seconds to readable format"""
        if seconds <= 0:
            return "Unknown"
        
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        seconds = seconds % 60
        
        if hours > 0:
            return f"{hours}h {minutes}m {seconds}s"
        elif minutes > 0:
            return f"{minutes}m {seconds}s"
        else:
            return f"{seconds}s"

# Global instance
export_service = ExportService()
