"""
Notes API Routes for Document Management and AI-Assisted Note Taking
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Body
from typing import List, Optional
from datetime import datetime
from app.models import (
    Document, DocumentCreate, DocumentUpdate, DocumentChatMessage,
    UserInDB, DocumentStatus
)
from app.auth import get_current_active_user
from app.database import get_database
from app.ai_service import ai_service
from bson import ObjectId
import logging
import os
import shutil
import docx
import PyPDF2
from io import BytesIO

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/notes", tags=["notes"])

# Upload directory for documents
UPLOAD_DIR = "static/uploads/documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@router.get("/documents", response_model=List[Document])
async def get_user_documents(
    search: Optional[str] = None,
    status: Optional[DocumentStatus] = None,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get all documents for the current user"""
    try:
        # Build query
        query = {"user_id": ObjectId(current_user.id)}
        
        if status:
            query["status"] = status
        
        # Get documents
        cursor = db.documents.find(query).sort("updated_at", -1)
        documents = await cursor.to_list(length=None)
        
        # Convert ObjectIds to strings and add id field
        for doc in documents:
            doc["id"] = str(doc["_id"])
        
        # Filter by search if provided
        if search:
            search_lower = search.lower()
            documents = [
                doc for doc in documents 
                if search_lower in doc.get("title", "").lower() or 
                   search_lower in doc.get("content", "").lower()
            ]
        
        return [Document(**doc) for doc in documents]
        
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch documents"
        )

@router.post("/documents", response_model=Document)
async def create_document(
    document: DocumentCreate,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Create a new document"""
    try:
        document_dict = {
            "user_id": ObjectId(current_user.id),
            "title": document.title,
            "content": document.content,
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(document_dict)
        document_dict["_id"] = result.inserted_id
        
        # Convert ObjectId to string for response
        document_dict["id"] = str(result.inserted_id)
        
        logger.info(f"Created document {result.inserted_id} for user {current_user.id}")
        return Document(**document_dict)
        
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create document"
        )

@router.get("/debug/test")
async def debug_test(
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Debug endpoint to test document creation and retrieval"""
    try:
        # Create a test document
        test_doc = {
            "user_id": ObjectId(current_user.id),
            "title": "Debug Test Document",
            "content": "This is a test document for debugging",
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(test_doc)
        test_doc["_id"] = result.inserted_id
        test_doc["id"] = str(result.inserted_id)
        
        logger.info(f"Created test document: {test_doc}")
        
        # Try to retrieve it
        retrieved = await db.documents.find_one({"_id": result.inserted_id})
        retrieved["id"] = str(retrieved["_id"])
        
        logger.info(f"Retrieved test document: {retrieved}")
        
        # Clean up
        await db.documents.delete_one({"_id": result.inserted_id})
        
        return {
            "created": test_doc,
            "retrieved": retrieved,
            "id_string": str(result.inserted_id),
            "id_valid": ObjectId.is_valid(str(result.inserted_id))
        }
        
    except Exception as e:
        logger.error(f"Debug test failed: {e}")
        return {"error": str(e)}

@router.get("/documents/{document_id}", response_model=Document)
async def get_document(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get a specific document"""
    logger.info(f"Getting document with ID: {document_id}")
    
    try:
        document_object_id = ObjectId(document_id)
        logger.info(f"Converted to ObjectId: {document_object_id}")
    except Exception as e:
        logger.error(f"Invalid document ID format: {document_id}, error: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        logger.warning(f"Document not found: {document_id} for user {current_user.id}")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Convert ObjectId to string for response
    document["id"] = str(document["_id"])
    logger.info(f"Returning document: {document['title']}")
    
    return Document(**document)

@router.put("/documents/{document_id}", response_model=Document)
async def update_document(
    document_id: str,
    document_update: DocumentUpdate,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Update a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Check if document exists and belongs to user
    existing_doc = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not existing_doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Build update data
    update_data = {"updated_at": datetime.utcnow()}
    if document_update.title is not None:
        update_data["title"] = document_update.title
    if document_update.content is not None:
        update_data["content"] = document_update.content
    if document_update.status is not None:
        update_data["status"] = document_update.status
    
    # Update document
    await db.documents.update_one(
        {"_id": document_object_id},
        {"$set": update_data}
    )
    
    # Return updated document
    updated_doc = await db.documents.find_one({"_id": document_object_id})
    updated_doc["id"] = str(updated_doc["_id"])
    return Document(**updated_doc)

@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Delete a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Check if document exists and belongs to user
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Delete associated file if exists
    if document.get("file_url"):
        try:
            if os.path.exists(document["file_url"]):
                os.remove(document["file_url"])
        except Exception as e:
            logger.warning(f"Failed to delete file {document['file_url']}: {e}")
    
    # Delete document and chat history
    await db.documents.delete_one({"_id": document_object_id})
    await db.document_chat_messages.delete_many({"document_id": document_object_id})
    
    return {"message": "Document deleted successfully"}

@router.post("/documents/upload", response_model=Document)
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Upload a document file and extract content"""
    try:
        # Validate file type
        allowed_extensions = ['.txt', '.docx', '.doc', '.pdf']
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only TXT, DOCX, DOC, and PDF files are allowed"
            )
        
        # Save file
        file_path = os.path.join(UPLOAD_DIR, f"{ObjectId()}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract content based on file type
        content = ""
        try:
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_extension in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file_extension == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = '\n'.join([page.extract_text() for page in pdf_reader.pages])
        except Exception as e:
            logger.warning(f"Failed to extract content from {file.filename}: {e}")
            content = f"Content extraction failed for {file.filename}. You can edit this document manually."
        
        # Use provided title or filename
        document_title = title or os.path.splitext(file.filename)[0]
        
        # Create document
        document_dict = {
            "user_id": ObjectId(current_user.id),
            "title": document_title,
            "content": content,
            "file_url": file_path,
            "file_name": file.filename,
            "file_size": os.path.getsize(file_path),
            "status": DocumentStatus.DRAFT,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.documents.insert_one(document_dict)
        document_dict["_id"] = result.inserted_id
        
        # Convert ObjectId to string for response
        document_dict["id"] = str(result.inserted_id)
        
        logger.info(f"Uploaded document {result.inserted_id} for user {current_user.id}")
        return Document(**document_dict)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to upload document"
        )

@router.post("/documents/{document_id}/chat")
async def chat_with_document(
    document_id: str,
    message: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Chat with AI about the document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Get recent chat history
        chat_history = await db.document_chat_messages.find({
            "document_id": document_object_id
        }).sort("timestamp", -1).limit(10).to_list(length=None)
        
        # Reverse to get chronological order
        chat_history.reverse()
        
        # Generate AI response
        ai_response = await ai_service.chat_with_document(
            content=document.get("content", ""),
            document_title=document.get("title", ""),
            user_message=message,
            chat_history=chat_history
        )
        
        # Save chat message
        chat_message = {
            "document_id": document_object_id,
            "user_id": ObjectId(current_user.id),
            "message": message,
            "response": ai_response,
            "timestamp": datetime.utcnow()
        }
        
        result = await db.document_chat_messages.insert_one(chat_message)
        chat_message["_id"] = result.inserted_id
        chat_message["id"] = str(result.inserted_id)
        
        return {
            "message": message,
            "response": ai_response,
            "timestamp": chat_message["timestamp"]
        }
        
    except Exception as e:
        logger.error(f"Error in document chat: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process chat message"
        )

@router.post("/documents/{document_id}/generate-notes")
async def generate_notes(
    document_id: str,
    prompt: str = Body(..., embed=True),
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Generate structured notes from document content"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Get document
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Generate notes using AI
        notes = await ai_service.generate_notes_from_document(
            content=document.get("content", ""),
            document_title=document.get("title", ""),
            user_prompt=prompt
        )
        
        return {
            "notes": notes,
            "prompt": prompt
        }
        
    except Exception as e:
        logger.error(f"Error generating notes: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate notes"
        )

@router.get("/documents/{document_id}/chat-history")
async def get_chat_history(
    document_id: str,
    current_user: UserInDB = Depends(get_current_active_user),
    db = Depends(get_database)
):
    """Get chat history for a document"""
    try:
        document_object_id = ObjectId(document_id)
    except:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID"
        )
    
    # Verify document ownership
    document = await db.documents.find_one({
        "_id": document_object_id,
        "user_id": ObjectId(current_user.id)
    })
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Get chat history
    chat_history = await db.document_chat_messages.find({
        "document_id": document_object_id
    }).sort("timestamp", 1).to_list(length=None)
    
    # Convert ObjectIds to strings
    for chat in chat_history:
        chat["id"] = str(chat["_id"])
    
    return {"chat_history": chat_history}